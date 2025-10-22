from abc import ABC, abstractmethod
from typing import Optional
import io
import logging
import PyPDF2
from docx import Document
import pandas as pd
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


class BaseFileExtractor(ABC):
    @abstractmethod
    async def extract(self, content: bytes, filename: Optional[str] = None) -> str:
        pass


class TXTTextExtractor(BaseFileExtractor):
    async def extract(
        self,
        content: bytes,
        filename: Optional[str] = None,
        chunk_size=1000,
        overlap=200,
    ) -> list:
        try:
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    text = content.decode("latin-1")
                except UnicodeDecodeError:
                    text = content.decode("utf-8", errors="ignore")
            text = text.replace("\r\n", "\n").replace("\r", "\n")
            return smart_chunk_text(text, chunk_size=chunk_size, overlap=overlap)
        except Exception:
            return []


class PDFTextExtractor(BaseFileExtractor):
    async def extract(
        self,
        content: bytes,
        filename: Optional[str] = None,
        chunk_size=1000,
        overlap=200,
    ) -> list:
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        page_text = page_text.replace("\r\n", "\n").replace("\r", "\n")
                        page_text = f"[Page {page_num + 1}]\n{page_text}"
                        text_parts.append(page_text)
                except Exception:
                    continue
            full_text = "\n".join(text_parts)
            return smart_chunk_text(full_text, chunk_size=chunk_size, overlap=overlap)
        except Exception as e:
            return []


def extract_tables_from_docx(doc):
    tables = []
    for table in doc.tables:
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)
        if rows:
            # Format as markdown table if possible
            header = rows[0]
            table_md = ["| " + " | ".join(header) + " |", "|" + "---|" * len(header)]
            for row in rows[1:]:
                table_md.append("| " + " | ".join(row) + " |")
            tables.append("\n".join(table_md))
    return tables


class DOCXTextExtractor(BaseFileExtractor):
    async def extract(
        self,
        content: bytes,
        filename: Optional[str] = None,
        chunk_size=1000,
        overlap=200,
    ) -> list:
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            text_parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            table_texts = extract_tables_from_docx(doc)
            # Separate narrative and tables with clear markers
            all_parts = []
            if text_parts:
                all_parts.append("\n\n".join(text_parts))
            for i, table in enumerate(table_texts):
                all_parts.append(f"\n[Table {i+1}]\n{table}")
            full_text = "\n\n".join(all_parts)
            return smart_chunk_text(full_text, chunk_size=chunk_size, overlap=overlap)
        except Exception as e:
            return []


class DOCLegacyTextExtractor(BaseFileExtractor):
    async def extract(
        self,
        content: bytes,
        filename: Optional[str] = None,
        chunk_size=1000,
        overlap=200,
    ) -> list:
        try:
            from langchain_community.document_loaders import (
                UnstructuredWordDocumentLoader,
            )
            import tempfile
            from utils.file_operations.file_manager import FileManager as FileUtils
            import gc

            temp_file_path = None
            try:
                with tempfile.NamedTemporaryFile(
                    suffix=".doc", delete=False
                ) as temp_file:
                    temp_file.write(content)
                    temp_file.flush()
                    temp_file_path = temp_file.name

                # Load documents
                loader = UnstructuredWordDocumentLoader(temp_file_path)
                documents = loader.load()
                text_parts = [
                    doc.page_content.strip()
                    for doc in documents
                    if doc.page_content and doc.page_content.strip()
                ]
                full_text = "\n\n".join(text_parts)
                result = smart_chunk_text(
                    full_text, chunk_size=chunk_size, overlap=overlap
                )

                # Clean up references to help with file release
                del documents
                del loader
                gc.collect()

                return result
            finally:
                # Ensure cleanup happens even if processing fails
                if temp_file_path:
                    try:
                        FileUtils.safe_delete_file(
                            temp_file_path, retry_count=5, retry_delay=0.2
                        )
                    except Exception as e:
                        logger.warning(
                            f"Failed to delete temporary file {temp_file_path}: {e}"
                        )
        except ImportError:
            return ["[UnstructuredWordDocumentLoader not available]"]
        except Exception as e:
            return []


def smart_chunk_rows(rows, chunk_size=1000, header=None):
    """
    Chunk a list of rows (strings) into groups, ensuring no row is split across chunks.
    Each chunk is a string of joined rows, up to chunk_size characters.
    Optionally, prepend a header to each chunk.
    Splits chunks at rows that are all NaN (empty or whitespace-only).
    """
    chunks = []
    current_chunk = []
    current_length = 0
    if header:
        header_length = len(header) + 1  # +1 for newline
    else:
        header_length = 0

    for i, row in enumerate(rows):
        row_length = len(row) + 1  # +1 for newline

        # Check if this row is all NaN (empty or whitespace-only)
        is_nan_row = (
            not row.strip()
            or row.strip() == "nan"
            or all(cell.strip() == "nan" for cell in row.split(" | "))
        )

        # If we have a current chunk and encounter a NaN row, start a new chunk
        if is_nan_row and current_chunk:
            chunk = (
                "\n".join([header] + current_chunk)
                if header
                else "\n".join(current_chunk)
            )
            chunks.append(chunk)
            current_chunk = []
            current_length = 0
            continue  # Skip adding the NaN row to chunks

        # Check if adding this row would exceed chunk size
        if current_length + row_length + header_length > chunk_size and current_chunk:
            # Start a new chunk
            chunk = (
                "\n".join([header] + current_chunk)
                if header
                else "\n".join(current_chunk)
            )
            chunks.append(chunk)
            current_chunk = []
            current_length = 0

        # Only add non-NaN rows to chunks
        if not is_nan_row:
            current_chunk.append(row)
            current_length += row_length

    # Add the last chunk if it exists
    if current_chunk:
        chunk = (
            "\n".join([header] + current_chunk) if header else "\n".join(current_chunk)
        )
        chunks.append(chunk)

    return chunks


class CSVTextExtractor(BaseFileExtractor):
    async def extract(
        self, content: bytes, filename: Optional[str] = None, chunk_size=1000
    ) -> list:
        try:
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    csv_text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                csv_text = content.decode("utf-8", errors="ignore")
            csv_file = io.StringIO(csv_text)
            df = pd.read_csv(csv_file)
            if df.empty:
                return []
            headers = " | ".join(df.columns)
            header = f"Columns: {headers}"

            # Process rows, handling NaN values properly
            rows = []
            for _, row in df.iterrows():
                # Convert each cell to string, handling NaN values
                row_cells = []
                for cell in row:
                    if pd.isna(cell):
                        row_cells.append("nan")
                    else:
                        row_cells.append(str(cell))
                row_str = " | ".join(row_cells)
                rows.append(row_str)

            return smart_chunk_rows(rows, chunk_size=chunk_size, header=header)
        except Exception as e:
            return []


class XLSXTextExtractor(BaseFileExtractor):
    async def extract(
        self, content: bytes, filename: Optional[str] = None, chunk_size=1000
    ) -> list:
        try:
            excel_file = io.BytesIO(content)
            # Use openpyxl with data_only=True so we get cached evaluated values instead of formulas
            wb = load_workbook(excel_file, data_only=True)
            all_chunks = []
            for ws in wb.worksheets:
                sheet_name = ws.title
                values = list(ws.iter_rows(values_only=True))
                if not values:
                    continue
                # Derive headers from first non-empty row; fallback to generic names
                header_row = None
                for r in values:
                    if any(cell is not None and str(cell).strip() != "" for cell in r):
                        header_row = r
                        break
                if header_row is None:
                    continue
                headers = [
                    (str(h).strip() if h is not None and str(h).strip() != "" else f"col_{i+1}")
                    for i, h in enumerate(header_row)
                ]
                header = f"Sheet: {sheet_name}\nColumns: {' | '.join(headers)}"

                # Build rows starting after the header row
                rows = []
                header_found = False
                for r in values:
                    if not header_found:
                        # skip until we've passed the header_row instance
                        if r is header_row:
                            header_found = True
                        continue
                    row_cells = []
                    empty_count = 0
                    for cell in r[: len(headers)]:
                        if cell is None:
                            row_cells.append("nan")
                            empty_count += 1
                        else:
                            # Format numbers to match Excel display (round to 2 decimal places)
                            if isinstance(cell, (int, float)):
                                if isinstance(cell, float) and cell.is_integer():
                                    row_cells.append(str(int(cell)))
                                else:
                                    row_cells.append(f"{cell:.2f}")
                            else:
                                row_cells.append(str(cell))
                    # Skip completely empty rows
                    if empty_count == len(headers):
                        rows.append("nan")  # marker to allow chunk splitter to break
                    else:
                        rows.append(" | ".join(row_cells))

                # Keep each sheet as one complete chunk instead of splitting by size
                # Filter out empty rows (marked as "nan")
                non_empty_rows = [row for row in rows if row != "nan"]
                
                if non_empty_rows:
                    # Create one chunk per sheet with all its data
                    sheet_content = "\n".join([header] + non_empty_rows)
                    all_chunks.append(sheet_content)
            return all_chunks
        except Exception as e:
            return []


class URLTextExtractor(BaseFileExtractor):
    async def extract(self, content: bytes, filename: Optional[str] = None) -> str:
        return "[URL extraction not implemented in extractor]"


def smart_chunk_text(text, chunk_size=1000, overlap=200):
    """
    Chunk narrative text, avoiding splitting in the middle of sentences.
    Delegates to TextUtils.chunk_text for unified chunking logic.
    """
    from utils.text_processing.text_utils import TextUtils

    return TextUtils.chunk_text(text, chunk_size=chunk_size, overlap=overlap)
